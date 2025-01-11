from docx import Document
from docx.shared import Pt

# создание документа
doc = Document()

# открытие документа
# document = Document('/path/to/document.docx')
# задаем стиль текста по умолчанию
style = doc.styles['Normal']
# название шрифта
style.font.name = 'Arial'
# размер шрифта
style.font.size = Pt(14)
doc.add_paragraph('Текст документа')

head = doc.add_heading('Основы работы с файлами Microsoft Word на Python.')
from docx.enum.text import WD_ALIGN_PARAGRAPH
# выравнивание посередине
head.alignment = WD_ALIGN_PARAGRAPH.CENTER


run = doc.add_paragraph().add_run('Заголовок, размером 24 pt.')
# размер шрифта
run.font.size = Pt(24)
run.bold = True
# doc.save('test.docx')

paragraph = doc.add_paragraph('Абзац содержит форматирование ')
paragraph.add_run('на уровне блока.').bold=True
doc.save("test_dock.docx")