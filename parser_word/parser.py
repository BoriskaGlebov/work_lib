from docx import Document

def format_text(text, line_length=60):
    """Форматирует текст, чтобы каждая строка не превышала заданную длину."""
    words = text.split()
    formatted_lines = []
    current_line = ""

    for word in words:
        if len(current_line) + len(word) + 1 <= line_length:
            current_line += " " + word if current_line else word
        else:
            formatted_lines.append(current_line)
            current_line = word

    if current_line:
        formatted_lines.append(current_line)

    return "\n".join(formatted_lines)

def parse_and_save(input_file, output_file):
    """Парсит документ Word и сохраняет отформатированный текст в новый файл."""
    doc = Document(input_file)
    full_text = []

    # Обработка параграфов
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text for cell in row.cells)  # Объединяем ячейки строки с разделителем
            full_text.append(row_text)

    combined_text = "\n".join(full_text)
    formatted_text = format_text(combined_text)

    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(formatted_text)

# Пример использования
input_file_path = 'file123.docx'  # Путь к исходному документу
output_file_path = 'name.txt'     # Путь к выходному файлу
parse_and_save(input_file_path, output_file_path)
