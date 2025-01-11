import os

from docx import Document


def format_text(text, max_length=60):
    """Форматирует текст, ограничивая длину строк до max_length символов."""
    words = text.split()
    formatted_lines = []
    current_line = []

    for word in words:
        # Проверяем, если добавление слова превышает максимальную длину
        if len(' '.join(current_line + [word])) <= max_length:
            current_line.append(word)
        else:
            # Сохраняем текущую строку и начинаем новую
            formatted_lines.append(' '.join(current_line))
            current_line = [word]

    # Добавляем последнюю строку, если она не пустая
    if current_line:
        formatted_lines.append(' '.join(current_line))

    return '\n'.join(formatted_lines)


# Открываем документ
document = Document('file123.docx')

# Открываем файл для записи результатов
with open('name.txt', 'w', encoding='utf-8') as output_file:
    # Читаем верхний колонтитул
    header = document.sections[0].header
    output_file.write("Верхний колонтитул:\n")
    for paragraph in header.paragraphs:
        output_file.write(format_text(paragraph.text) + '\n')
    output_file.write('*' * 50 + '\n')

    # Читаем основной текст документа и таблицы
    output_file.write("\nСодержимое документа:\n")
    num_tables = 0
    for element in document.element.body:
        if element.tag.endswith('p'):  # Проверяем, является ли элемент абзацем
            text = element.text.strip()
            if text:  # Если абзац не пустой
                output_file.write(format_text(text) + '\n')
            else:  # Печатаем пустую строку для пустого абзаца
                output_file.write('\n')
        elif element.tag.endswith('tbl') and num_tables < len(document.tables):
            output_file.write('-' * 60 + '\n')
            table = document.tables[num_tables]  # Получаем таблицу по индексу
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        output_file.write(format_text(cell.text.strip()) + '\n')  # Форматируем содержимое ячейки
                output_file.write('\n')  # Переход на новую строку после каждой строки таблицы
            num_tables += 1
            output_file.write('-' * 60)
    output_file.write('*' * 50 + '\n')

    # Читаем нижний колонтитул
    footer = document.sections[0].footer
    output_file.write("\nНижний колонтитул:\n")
    for paragraph in footer.paragraphs:
        output_file.write(format_text(paragraph.text) + '\n')
if __name__ == '__main__':
    print(os.path.exists('file123.docx'))
    print(os.path.splitext('file123.docx')[0])
    num=input("Введи номер")
    print(os.path.splitext('file123.docx')[0]+'_'+num)